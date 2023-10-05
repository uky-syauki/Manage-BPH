from docx import Document
from docx.shared import Inches
import os

class CreadDocx:
    def __init__(self):
        self.__documen = Document()
    def tambahkan_judul(self, judul, tanggal):
        j = self.__documen.add_heading(judul, 0)
        j.add_run(tanggal.strftime("%Y/%m/%d")).italic = True
    def tambahkan_heading(self, heading, h):
        self.__documen.add_heading(heading, level=h)
    def tambahkan_quote(self, quote):
        self.__documen.add_paragraph(quote, style="Intense Quote")
    def tambahkan_paragraph(self, paragraph, bold=None):
        p = self.__documen.add_paragraph('')
        if bold is not None:
            p.add_run(paragraph).bold = True
        else:
            p.add_run(paragraph)
    def buat(self,nama):
        self.__documen.save(f'app/document/{nama}.docx')


class EditDocx:
    def __init__(self,path):
        self.__path = path
        self.__document = Document(path)
        print(len(self.__document.paragraphs))
    def cetak(self):
        for paragraph in self.__document.paragraphs:
            print(paragraph.text)
    def edit_dan_save(self):
        for paragraph in self.__document.paragraphs:
            paragraph.text = "Dirubah"
    def save(self):
        self.__document.save(self.__path)


class Humas:
    def ambil_list():
        return os.listdir('app/document/Humas')

        



